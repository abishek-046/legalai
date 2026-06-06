"""
OCR and document text extraction utilities

Extraction pipeline:
1. PDF (text-based)  → pdfplumber + PyMuPDF direct text
2. PDF (scanned)     → OCR.space API (cloud OCR, no system deps)
3. DOCX / DOC        → python-docx
4. Images            → OCR.space API → pytesseract fallback
"""

import io
import base64
import logging
import re
import httpx
from pathlib import Path

logger = logging.getLogger(__name__)

PDF_MIN_CHARS = 80
PAGE_MIN_CHARS = 20


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ─── OCR.space API ────────────────────────────────────────────────────────────

async def _ocr_space_from_bytes(
    file_bytes: bytes,
    filename: str,
    api_key: str,
    is_pdf: bool = False,
) -> str:
    """
    Send file bytes to OCR.space API and return extracted text.
    Free tier key: 'helloworld' (25,000 req/month, max 1MB).
    Get a free key at: https://ocr.space/OCRAPI
    """
    try:
        # Encode file as base64
        b64 = base64.b64encode(file_bytes).decode("utf-8")
        ext = Path(filename).suffix.lower().lstrip(".")
        mime_map = {
            "pdf": "application/pdf",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "tif": "image/tiff",
            "webp": "image/webp",
        }
        mime = mime_map.get(ext, "application/octet-stream")
        data_uri = f"data:{mime};base64,{b64}"

        payload = {
            "apikey": api_key,
            "base64Image": data_uri,
            "language": "eng",
            "isOverlayRequired": False,
            "detectOrientation": True,
            "scale": True,
            "OCREngine": 2,  # Engine 2 is better for printed text
        }

        if is_pdf:
            payload["isTable"] = False

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.ocr.space/parse/image",
                data=payload,
            )

        result = response.json()

        if result.get("IsErroredOnProcessing"):
            err_msg = result.get("ErrorMessage", ["Unknown OCR error"])
            logger.warning(f"OCR.space error: {err_msg}")
            return ""

        parsed = result.get("ParsedResults", [])
        if not parsed:
            return ""

        texts = [p.get("ParsedText", "") for p in parsed if p.get("ParsedText")]
        combined = "\n".join(texts)
        logger.info(f"OCR.space extracted {len(combined)} chars from {filename}")
        return combined

    except Exception as e:
        logger.error(f"OCR.space API failed: {e}")
        return ""


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _try_pdfplumber(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if len(t.strip()) >= PAGE_MIN_CHARS:
                    parts.append(t.strip())
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning(f"pdfplumber: {e}")
        return ""


def _try_pymupdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
        parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            # Try multiple extraction modes
            t = page.get_text("text") or ""
            if len(t.strip()) < PAGE_MIN_CHARS:
                # Try words mode
                words = page.get_text("words")
                if words:
                    t = " ".join(w[4] for w in words if isinstance(w[4], str))
            if len(t.strip()) >= PAGE_MIN_CHARS:
                parts.append(t.strip())
        doc.close()
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning(f"PyMuPDF: {e}")
        return ""


def extract_from_pdf_sync(file_bytes: bytes) -> str:
    """Synchronous PDF extraction (text-based PDFs only)."""
    text = _try_pdfplumber(file_bytes)
    if len(text.strip()) >= PDF_MIN_CHARS:
        logger.info(f"pdfplumber OK: {len(text)} chars")
        return clean_text(text)

    text2 = _try_pymupdf_text(file_bytes)
    if len(text2.strip()) >= PDF_MIN_CHARS:
        logger.info(f"PyMuPDF OK: {len(text2)} chars")
        return clean_text(text2)

    return clean_text(text or text2)


async def extract_from_pdf(file_bytes: bytes, filename: str = "doc.pdf") -> str:
    """
    Full PDF extraction pipeline - optimized for speed on Render free tier.
    Only processes first 3 pages to stay under 30s timeout.
    """
    from config import settings

    logger.info(f"PDF: {filename} ({len(file_bytes)} bytes)")

    # For large PDFs, trim to first 3 pages before processing
    if len(file_bytes) > 500_000:
        logger.info("Large PDF detected — trimming to first 3 pages for speed")
        file_bytes = _trim_pdf(file_bytes, max_pages=3)

    # Step 1: direct text extraction (fast)
    text = extract_from_pdf_sync(file_bytes)
    if len(text.strip()) >= PDF_MIN_CHARS:
        return text

    # Step 2: OCR.space for scanned PDF
    logger.info("Direct extraction insufficient — using OCR.space")

    ocr_text = await _ocr_space_from_bytes(
        file_bytes, filename, settings.OCR_SPACE_API_KEY, is_pdf=True
    )

    if len(ocr_text.strip()) >= 20:
        logger.info(f"OCR.space OK: {len(ocr_text)} chars")
        return clean_text(ocr_text)

    logger.warning(f"All PDF methods limited. Best: {len(text)} chars")
    return text


def _trim_pdf(file_bytes: bytes, max_pages: int = 5) -> bytes:
    """Return a PDF with only the first N pages (reduces size for OCR API)."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if len(doc) <= max_pages:
            doc.close()
            return file_bytes

        new_doc = fitz.open()
        for i in range(min(max_pages, len(doc))):
            new_doc.insert_pdf(doc, from_page=i, to_page=i)

        result = new_doc.tobytes()
        new_doc.close()
        doc.close()
        logger.info(f"PDF trimmed to {max_pages} pages")
        return result
    except Exception:
        return file_bytes


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def extract_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text.strip() for c in row.cells if c.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        result = clean_text("\n".join(parts))
        logger.info(f"DOCX: {len(result)} chars")
        return result
    except ImportError:
        raise RuntimeError("DOCX library not available")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


# ─── Images ───────────────────────────────────────────────────────────────────

async def extract_from_image(file_bytes: bytes, filename: str = "img.jpg") -> str:
    """Extract text from image using OCR.space API."""
    from config import settings

    logger.info(f"Image OCR: {filename} ({len(file_bytes)} bytes)")

    # Try OCR.space
    text = await _ocr_space_from_bytes(
        file_bytes, filename, settings.OCR_SPACE_API_KEY
    )
    if text.strip():
        return clean_text(text)

    # Fallback: pytesseract
    try:
        import pytesseract
        from PIL import Image
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")
        t = pytesseract.image_to_string(image, lang="eng")
        result = clean_text(t)
        if result:
            return result
    except Exception as e:
        logger.warning(f"pytesseract: {e}")

    raise RuntimeError(
        "Could not extract text from this image. "
        "Please ensure the image contains clear, readable text."
    )


# ─── Entry Point ──────────────────────────────────────────────────────────────

async def extract_text_async(file_bytes: bytes, filename: str) -> str:
    """Async entry point for text extraction."""
    ext = Path(filename).suffix.lower()
    if not file_bytes:
        raise ValueError("File is empty")
    if ext == ".pdf":
        return await extract_from_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"):
        return await extract_from_image(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            "Supported: PDF, DOCX, DOC, PNG, JPG, JPEG, WEBP, BMP, TIFF"
        )


# Keep sync version for backward compat
def extract_text(file_bytes: bytes, filename: str) -> str:
    """Sync wrapper — use extract_text_async for new code."""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, extract_text_async(file_bytes, filename))
                return future.result(timeout=90)
        else:
            return loop.run_until_complete(extract_text_async(file_bytes, filename))
    except Exception as e:
        raise RuntimeError(str(e))


def get_document_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return {
        ".pdf": "PDF", ".docx": "DOCX", ".doc": "DOC",
        ".png": "Image (PNG)", ".jpg": "Image (JPG)", ".jpeg": "Image (JPEG)",
        ".webp": "Image (WEBP)", ".bmp": "Image (BMP)", ".tiff": "Image (TIFF)",
    }.get(ext, "Unknown")
